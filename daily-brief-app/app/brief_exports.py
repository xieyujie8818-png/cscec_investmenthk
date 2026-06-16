"""Export daily brief to Word (.docx) and styled HTML."""

from __future__ import annotations

import subprocess
from copy import deepcopy
from datetime import date
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.models import DailyReport, ScoredArticle
from app.scoring import extract_body_paragraphs

TEMPLATES_DIR = Path(__file__).resolve().parent.parent / "templates"
BRIEF_TEMPLATE_DOC = TEMPLATES_DIR / "brief-template.doc"
BRIEF_TEMPLATE_PATH = TEMPLATES_DIR / "brief-template.docx"

# Prototype paragraph indices from brief-template.docx (663期 reference layout)
COVER_END = 6
TOC_HEADING = 6
TOC_SECTION = 7
TOC_ENTRY = 8
TOC_SUMMARY = 9
BODY_TITLE_FIRST = 32
BODY_TITLE_NEXT = 47
BODY_SOURCE = 33
BODY_PARAGRAPH = 34

TOC_PAGEREF_PLACEHOLDER = "1"


def _parse_pub_date(value: str, fallback: date) -> date:
    if not value:
        return fallback
    try:
        return date.fromisoformat(value[:10])
    except ValueError:
        return fallback


def _item_dateline(item: ScoredArticle, fallback: date) -> str:
    pub = _parse_pub_date(item.publish_date, fallback)
    return f"【本報{pub.month}月{pub.day}日消息】"


def _item_paragraphs(item: ScoredArticle) -> list[str]:
    if item.body:
        return extract_body_paragraphs(item.body)
    if item.paragraphs:
        return list(item.paragraphs)
    if item.body_paragraphs:
        return list(item.body_paragraphs)
    return []


def _add_field_to_run(run, instr_text: str, placeholder: str) -> None:
    r = run._r
    for child in list(r):
        r.remove(child)

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instr_text

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = placeholder
    fld_sep.append(text)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)


def _enable_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)


def _clear_document_body(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _convert_doc_to_docx(doc_path: Path, docx_path: Path) -> None:
    ps = f"""
$doc = "{doc_path.resolve()}"
$docx = "{docx_path.resolve()}"
$word = New-Object -ComObject Word.Application
$word.Visible = $false
try {{
  $d = $word.Documents.Open($doc)
  $d.SaveAs2([ref]$docx, [ref]16)
  $d.Close([ref]$false)
}} finally {{
  $word.Quit()
  [System.Runtime.InteropServices.Marshal]::ReleaseComObject($word) | Out-Null
}}
"""
    subprocess.run(
        ["powershell", "-NoProfile", "-Command", ps],
        check=True,
        capture_output=True,
        text=True,
    )


def _ensure_brief_template_docx() -> Path:
    if BRIEF_TEMPLATE_PATH.exists():
        return BRIEF_TEMPLATE_PATH
    if BRIEF_TEMPLATE_DOC.exists():
        _convert_doc_to_docx(BRIEF_TEMPLATE_DOC, BRIEF_TEMPLATE_PATH)
        return BRIEF_TEMPLATE_PATH
    raise FileNotFoundError(
        f"Word 模板不存在，請放置 {BRIEF_TEMPLATE_PATH} 或 {BRIEF_TEMPLATE_DOC}"
    )


def _append_cloned_paragraph(doc: Document, src_paragraph) -> None:
    doc.element.body.append(deepcopy(src_paragraph._p))


def _replace_paragraph_plain_text(paragraph, text: str) -> None:
    t_nodes = list(paragraph._p.iter(qn("w:t")))
    if not t_nodes:
        paragraph.add_run(text)
        return
    t_nodes[0].text = text
    for node in t_nodes[1:]:
        node.text = ""


def _replace_heading_title(paragraph, title: str) -> None:
    runs = list(paragraph.runs)
    start_idx = 0
    if runs and runs[0]._r.find(qn("w:br")) is not None:
        start_idx = 1
    elif len(runs) >= 4 and not (runs[0].text or "").strip():
        start_idx = 1

    content_runs = runs[start_idx:]
    clean = title.strip()
    if len(content_runs) >= 3:
        if " " in clean:
            part1, part2 = clean.split(" ", 1)
        else:
            part1, part2 = clean, ""
        content_runs[0].text = part1
        content_runs[1].text = " " if part2 else ""
        content_runs[2].text = part2
        for run in content_runs[3:]:
            if run.text is not None:
                run.text = ""
        return
    _replace_paragraph_plain_text(paragraph, clean)


def _clear_bookmarks(paragraph) -> None:
    for tag in ("w:bookmarkStart", "w:bookmarkEnd"):
        for element in list(paragraph._p.iter(qn(tag))):
            element.getparent().remove(element)


def _add_bookmark(paragraph, bookmark_name: str, bookmark_id: int) -> None:
    _clear_bookmarks(paragraph)

    bookmark_start = OxmlElement("w:bookmarkStart")
    bookmark_start.set(qn("w:id"), str(bookmark_id))
    bookmark_start.set(qn("w:name"), bookmark_name)

    bookmark_end = OxmlElement("w:bookmarkEnd")
    bookmark_end.set(qn("w:id"), str(bookmark_id))

    insert_idx = 0
    for idx, child in enumerate(paragraph._p):
        if child.tag == qn("w:r") and child.find(qn("w:br")) is not None:
            insert_idx = idx + 1
            break

    paragraph._p.insert(insert_idx, bookmark_start)
    paragraph._p.append(bookmark_end)


def _set_toc_pageref(paragraph, bookmark_name: str) -> None:
    for instr in paragraph._p.iter(qn("w:instrText")):
        if instr.text and "PAGEREF" in instr.text:
            instr.text = f" PAGEREF {bookmark_name} \\h "
            return

    runs = list(paragraph.runs)
    if len(runs) >= 5:
        _add_field_to_run(
            runs[4],
            f"PAGEREF {bookmark_name} \\h",
            TOC_PAGEREF_PLACEHOLDER,
        )


def _set_toc_line(paragraph, title: str, bookmark_name: str) -> None:
    runs = list(paragraph.runs)
    clean_title = title.strip()
    if len(runs) >= 9:
        if " " in clean_title:
            part1, part2 = clean_title.split(" ", 1)
        else:
            part1, part2 = clean_title, ""
        runs[0].text = part1
        runs[1].text = " " if part2 else ""
        runs[2].text = part2
        for run in runs[3:]:
            if run.text and run.text != "\t":
                run.text = ""
        if len(runs) > 8:
            runs[8].text = TOC_PAGEREF_PLACEHOLDER
        _set_toc_pageref(paragraph, bookmark_name)
        return

    paragraph.clear()
    paragraph.add_run(clean_title)
    paragraph.add_run("\t")
    _add_field_to_run(
        paragraph.add_run(),
        f"PAGEREF {bookmark_name} \\h",
        TOC_PAGEREF_PLACEHOLDER,
    )


def _add_cover_page(doc: Document, ref: Document, report: DailyReport) -> None:
    for idx in range(COVER_END):
        _append_cloned_paragraph(doc, ref.paragraphs[idx])

    cover_date = doc.paragraphs[2]
    _replace_paragraph_plain_text(
        cover_date,
        f"{report.weekday_label}  海外與投資部 編制",
    )


def _add_toc_section(
    doc: Document,
    ref: Document,
    report: DailyReport,
    bookmark_id_start: int,
) -> list[tuple[ScoredArticle, str, int]]:
    bookmark_id = bookmark_id_start
    bookmarks: list[tuple[ScoredArticle, str, int]] = []

    _append_cloned_paragraph(doc, ref.paragraphs[TOC_HEADING])

    for section_name, items in report.sections.items():
        if not items:
            continue

        _append_cloned_paragraph(doc, ref.paragraphs[TOC_SECTION])
        _replace_paragraph_plain_text(doc.paragraphs[-1], f"【{section_name}】")

        for item in items:
            bookmark_name = f"article_{item.news_id or bookmark_id}"
            _append_cloned_paragraph(doc, ref.paragraphs[TOC_ENTRY])
            toc_para = doc.paragraphs[-1]
            _set_toc_line(toc_para, item.title, bookmark_name)

            if item.catalog_summary:
                _append_cloned_paragraph(doc, ref.paragraphs[TOC_SUMMARY])
                _replace_paragraph_plain_text(
                    doc.paragraphs[-1],
                    f"【{item.catalog_summary}】",
                )

            bookmarks.append((item, bookmark_name, bookmark_id))
            bookmark_id += 1

    _append_cloned_paragraph(doc, ref.paragraphs[30])
    _append_cloned_paragraph(doc, ref.paragraphs[31])

    return bookmarks


def _add_body_content(
    doc: Document,
    ref: Document,
    bookmarks: list[tuple[ScoredArticle, str, int]],
    fallback: date,
) -> None:
    for article_index, (item, bookmark_name, bookmark_id) in enumerate(bookmarks):
        title_ref = (
            ref.paragraphs[BODY_TITLE_FIRST]
            if article_index == 0
            else ref.paragraphs[BODY_TITLE_NEXT]
        )
        _append_cloned_paragraph(doc, title_ref)

        title_para = doc.paragraphs[-1]
        _replace_heading_title(title_para, item.title)
        _add_bookmark(title_para, bookmark_name, bookmark_id)

        _append_cloned_paragraph(doc, ref.paragraphs[BODY_SOURCE])
        _replace_paragraph_plain_text(
            doc.paragraphs[-1],
            f"[來源：{item.source_name}]",
        )

        paragraphs = _item_paragraphs(item)
        if not paragraphs:
            _append_cloned_paragraph(doc, ref.paragraphs[BODY_PARAGRAPH])
            continue

        first = paragraphs[0]
        if first.startswith("【本報") and "日消息】" in first:
            body_text = first
        else:
            body_text = f"{_item_dateline(item, fallback)}{first}"

        _append_cloned_paragraph(doc, ref.paragraphs[BODY_PARAGRAPH])
        _replace_paragraph_plain_text(doc.paragraphs[-1], body_text)

        for para in paragraphs[1:]:
            _append_cloned_paragraph(doc, ref.paragraphs[BODY_PARAGRAPH])
            _replace_paragraph_plain_text(doc.paragraphs[-1], para)


def write_word_brief(report: DailyReport, out_path: Path, date_end: date | None = None) -> None:
    template_path = _ensure_brief_template_docx()
    ref = Document(str(template_path))
    doc = Document(str(template_path))
    _clear_document_body(doc)
    _enable_update_fields_on_open(doc)

    fallback = date_end
    if fallback is None:
        try:
            fallback = date.fromisoformat(report.date_end or report.date)
        except ValueError:
            fallback = date.today()

    _add_cover_page(doc, ref, report)
    bookmarks = _add_toc_section(doc, ref, report, bookmark_id_start=1)
    _add_body_content(doc, ref, bookmarks, fallback)

    doc.save(str(out_path))
